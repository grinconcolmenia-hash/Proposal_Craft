"""
ProposalCraft — Motor de Cotizaciones de Equipos / Alquiler
Diseñado para: listas de ítems con precio unitario, múltiples fechas de evento,
subtotal por fecha y total general.

Uso:
    from src.brand_loader import load_brand
    from src.quote_engine import generate_quote, QuoteData, QuoteItem, EventDate

    brand  = load_brand()
    quote  = QuoteData(
        cliente     = "UNEQUAL",
        nit_cliente = "NIT 901170356-3",
        subtitulo   = "Copa Mundial de Fútbol · Grupo K",
        fechas      = [
            EventDate("Partido 1", "Uzbekistán vs Colombia", "17 jun · 9:00 p.m."),
            ...
        ],
        items       = [
            QuoteItem("Pantalla 3×5",          3_750_000),
            ...
        ],
        forma_pago  = "Pago a 30 días",
        nota_pago   = "Precio aplica por cada fecha de evento.",
    )
    generate_quote(brand, quote)
"""

import os
from dataclasses import dataclass, field
from datetime import date, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.brand_loader import BrandConfig


# ══════════════════════════════════════════════════════════════════════════════
# ESTRUCTURAS DE DATOS
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class QuoteItem:
    """Una línea de equipo/ítem en la cotización."""
    descripcion: str
    valor: int                  # valor entero en COP (sin puntos)
    cantidad: int = 1           # si > 1 se muestra en la descripción
    nota: str = ""              # texto adicional pequeño bajo la descripción


@dataclass
class EventDate:
    """Una fecha/partido del evento."""
    label: str       # "Partido 1"
    nombre: str      # "Uzbekistán vs Colombia"
    detalle: str     # "17 jun · 9:00 p.m."


@dataclass
class QuoteData:
    cliente: str
    nit_cliente: str
    subtitulo: str
    items: list[QuoteItem]
    fechas: list[EventDate]

    forma_pago: str  = "Pago a 30 días"
    nota_pago: str   = ""
    incluye_extra: str = ""     # texto libre bajo 'Incluye' (opcional)

    fecha_emision: str | None  = None
    fecha_vigencia: str | None = None
    output_filename: str | None = None


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS (todos leen de brand, sin hardcoding)
# ══════════════════════════════════════════════════════════════════════════════

def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _cop(valor: int) -> str:
    """Formatea un entero como $ X.XXX.XXX"""
    return f"$ {valor:,.0f}".replace(",", ".")


def _set_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    tcPr.append(shd)


def _set_padding(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_col_w(cell, twips: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    bdr = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        bdr.append(b)
    tblPr.append(bdr)


def _tbl_width(table, twips: int):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    if line:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _run(para, text: str, brand: BrandConfig,
         bold=False, italic=False, size=9.5,
         color: RGBColor = None) -> None:
    if color is None:
        color = brand.color_body_text
    r = para.add_run(text)
    r.bold  = bold
    r.italic = italic
    r.font.name  = brand.font
    r.font.size  = Pt(size)
    r.font.color.rgb = color


def _para(doc, brand: BrandConfig,
          text="", bold=False, italic=False, size=9.5,
          color: RGBColor = None,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          before=0, after=60, line=None):
    if color is None:
        color = brand.color_body_text
    p = doc.add_paragraph()
    p.alignment = align
    _spacing(p, before=before, after=after, line=line)
    if text:
        _run(p, text, brand, bold=bold, italic=italic, size=size, color=color)
    return p


def _divider(doc, color_hex: str, sz=6, before=80, after=80):
    p = doc.add_paragraph()
    _spacing(p, before=before, after=after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _section_label(doc, brand: BrandConfig, text: str):
    _para(doc, brand, text.upper(), bold=True, size=8,
          color=brand.color_accent_1, before=0, after=28)


# ══════════════════════════════════════════════════════════════════════════════
# TABLA DE FECHAS
# ══════════════════════════════════════════════════════════════════════════════

def _build_fechas_table(doc, brand: BrandConfig, fechas: list[EventDate]):
    """3 columnas: label oscuro | nombre partido | detalle (hora/fecha)"""
    FULL  = 9360
    W_LBL = 1100   # "Partido N"
    W_NOM = 5200   # nombre del partido
    W_DET = 3060   # fecha/hora

    t = doc.add_table(rows=1 + len(fechas), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(t, FULL)
    _no_borders(t)

    # Header
    hdr_data = [("", W_LBL), ("PARTIDO", W_NOM), ("FECHA / HORA", W_DET)]
    for ci, (txt, w) in enumerate(hdr_data):
        c = t.rows[0].cells[ci]
        _set_bg(c, brand.color_primary_dark)
        _set_col_w(c, w)
        _set_padding(c, top=70, bottom=70)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before=0, after=0)
        if txt:
            _run(p, txt, brand, bold=True, size=7.5, color=brand.color_accent_1)

    # Filas
    for ri, fd in enumerate(fechas):
        row = t.rows[ri + 1]
        bg = brand.color_light_bg if ri % 2 == 0 else brand.color_white

        # col 0 — label
        c0 = row.cells[0]
        _set_bg(c0, brand.color_primary_dark)
        _set_col_w(c0, W_LBL)
        _set_padding(c0, top=80, bottom=80)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p0, before=0, after=0)
        _run(p0, fd.label, brand, bold=True, size=7.5, color=brand.color_accent_1)

        # col 1 — nombre partido
        c1 = row.cells[1]
        _set_bg(c1, bg)
        _set_col_w(c1, W_NOM)
        _set_padding(c1, top=80, bottom=80, left=140)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        _spacing(p1, before=0, after=0)
        _run(p1, fd.nombre, brand, bold=True, size=9.5, color=brand.color_body_text)

        # col 2 — detalle
        c2 = row.cells[2]
        _set_bg(c2, bg)
        _set_col_w(c2, W_DET)
        _set_padding(c2, top=80, bottom=80)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p2, before=0, after=0)
        _run(p2, fd.detalle, brand, bold=False, size=8.5, color=brand.color_mid_gray)


# ══════════════════════════════════════════════════════════════════════════════
# TABLA DE ITEMS (equipos / servicios)
# ══════════════════════════════════════════════════════════════════════════════

def _build_items_table(doc, brand: BrandConfig,
                       items: list[QuoteItem],
                       subtotal: int, n_fechas: int, total: int):
    """
    3 columnas: N° | Descripción | Valor
    Al final: fila SUBTOTAL/FECHA (resaltada) + fila TOTAL (más resaltada).
    """
    FULL  = 9360
    W_NUM = 700
    W_DES = 6760
    W_VAL = 1900

    n_items = len(items)
    n_rows  = 1 + n_items + 2   # header + items + subtotal + total

    t = doc.add_table(rows=n_rows, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(t, FULL)
    _no_borders(t)

    # ── Header ────────────────────────────────────────────────────────────────
    hdr_data = [("N°", W_NUM), ("DESCRIPCIÓN", W_DES), ("VALOR / FECHA", W_VAL)]
    for ci, (txt, w) in enumerate(hdr_data):
        c = t.rows[0].cells[ci]
        _set_bg(c, brand.color_primary_dark)
        _set_col_w(c, w)
        _set_padding(c, top=70, bottom=70)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        _spacing(p, before=0, after=0)
        _run(p, txt, brand, bold=True, size=7.5, color=brand.color_accent_1)

    # ── Ítems ─────────────────────────────────────────────────────────────────
    for ri, item in enumerate(items):
        row = t.rows[ri + 1]
        bg  = brand.color_light_bg if ri % 2 == 0 else brand.color_white

        c0 = row.cells[0]
        _set_bg(c0, brand.color_primary_dark)
        _set_col_w(c0, W_NUM)
        _set_padding(c0, top=80, bottom=80)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p0, before=0, after=0)
        _run(p0, f"{ri+1:02d}", brand, bold=True, size=8, color=brand.color_mid_gray)

        c1 = row.cells[1]
        _set_bg(c1, bg)
        _set_col_w(c1, W_DES)
        _set_padding(c1, top=80, bottom=item.nota and 40 or 80, left=140)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        _spacing(p1, before=0, after=0)
        _run(p1, item.descripcion, brand, bold=False, size=9.5, color=brand.color_body_text)
        if item.nota:
            _run(p1, f"\n{item.nota}", brand, italic=True, size=7.5, color=brand.color_mid_gray)

        c2 = row.cells[2]
        _set_bg(c2, bg)
        _set_col_w(c2, W_VAL)
        _set_padding(c2, top=80, bottom=80)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p2, before=0, after=0)
        _run(p2, _cop(item.valor), brand, bold=False, size=9.5, color=brand.color_body_text)

    # ── Fila SUBTOTAL POR FECHA ───────────────────────────────────────────────
    r_sub = t.rows[1 + n_items]

    c0 = r_sub.cells[0]
    _set_bg(c0, brand.color_accent_2)
    _set_col_w(c0, W_NUM)
    _set_padding(c0, top=90, bottom=90)
    _spacing(c0.paragraphs[0], before=0, after=0)

    c1 = r_sub.cells[1]
    _set_bg(c1, brand.color_accent_2)
    _set_col_w(c1, W_DES)
    _set_padding(c1, top=90, bottom=90, left=140)
    p1 = c1.paragraphs[0]
    _spacing(p1, before=0, after=0)
    _run(p1, f"SUBTOTAL POR FECHA  ×  {n_fechas} fecha{'s' if n_fechas != 1 else ''}",
         brand, bold=True, size=9, color=brand.color_accent_1)

    c2 = r_sub.cells[2]
    _set_bg(c2, brand.color_accent_2)
    _set_col_w(c2, W_VAL)
    _set_padding(c2, top=90, bottom=90)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _spacing(p2, before=0, after=0)
    _run(p2, _cop(subtotal), brand, bold=True, size=9, color=brand.color_accent_1)

    # ── Fila TOTAL GENERAL ────────────────────────────────────────────────────
    r_tot = t.rows[1 + n_items + 1]

    c0 = r_tot.cells[0]
    _set_bg(c0, brand.color_primary_dark)
    _set_col_w(c0, W_NUM)
    _set_padding(c0, top=100, bottom=100)
    _spacing(c0.paragraphs[0], before=0, after=0)

    c1 = r_tot.cells[1]
    _set_bg(c1, brand.color_primary_dark)
    _set_col_w(c1, W_DES)
    _set_padding(c1, top=100, bottom=100, left=140)
    p1 = c1.paragraphs[0]
    _spacing(p1, before=0, after=0)
    _run(p1, "TOTAL GENERAL", brand, bold=True, size=10.5, color=brand.color_accent_1)

    c2 = r_tot.cells[2]
    _set_bg(c2, brand.color_primary_dark)
    _set_col_w(c2, W_VAL)
    _set_padding(c2, top=100, bottom=100)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _spacing(p2, before=0, after=0)
    _run(p2, _cop(total), brand, bold=True, size=12, color=brand.color_accent_1)


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

def generate_quote(brand: BrandConfig, quote: QuoteData,
                   output_path: str | None = None) -> str:
    """
    Genera un documento Word (.docx) para cotización de equipos/alquiler.

    Returns:
        Ruta absoluta del archivo generado.
    """
    MESES = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]

    def fmt_date(d: date) -> str:
        return f"{d.day} de {MESES[d.month - 1]} de {d.year}"

    today          = date.today()
    fecha_emision  = quote.fecha_emision  or fmt_date(today)
    fecha_vigencia = quote.fecha_vigencia or fmt_date(today + timedelta(days=brand.validity_days))

    # Cálculos
    subtotal_fecha = sum(i.valor for i in quote.items)
    n_fechas       = len(quote.fechas)
    total_general  = subtotal_fecha * n_fechas

    # Ruta de salida
    if output_path is None:
        fname = quote.output_filename or f"Cotizacion_{quote.cliente}.docx"
        output_path = os.path.join(brand.output_dir, fname)

    # ── Documento ─────────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    FULL = 9360   # twips útiles

    # ── 1. LOGO ───────────────────────────────────────────────────────────────
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_logo, before=0, after=40)
    logo_path = brand.logo_dark_path or brand.logo_light_path
    if logo_path:
        p_logo.add_run().add_picture(logo_path, width=Cm(4))
    else:
        _run(p_logo, brand.company_name, brand,
             bold=True, size=16, color=brand.color_accent_1)

    _divider(doc, _rgb_hex(brand.color_accent_1), sz=12, before=20, after=50)

    # ── 2. ENCABEZADO ─────────────────────────────────────────────────────────
    _para(doc, brand, "COTIZACIÓN COMERCIAL", bold=True, size=7.5,
          color=brand.color_mid_gray, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=0, after=4)
    _para(doc, brand, quote.subtitulo, italic=True, size=9,
          color=brand.color_mid_gray, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=0, after=14)

    # Bloque cliente — tabla 2 cols: label | valor
    tcl = doc.add_table(rows=2, cols=2)
    tcl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_width(tcl, FULL)
    _no_borders(tcl)

    for ri, (lbl, val) in enumerate([
        ("Cliente",  quote.cliente),
        ("NIT",      quote.nit_cliente),
    ]):
        bg = brand.color_primary_dark
        c0 = tcl.rows[ri].cells[0]
        _set_bg(c0, bg)
        _set_col_w(c0, 1600)
        _set_padding(c0, top=60, bottom=60, left=100)
        _spacing(c0.paragraphs[0], before=0, after=0)
        _run(c0.paragraphs[0], lbl, brand, bold=True, size=8,
             color=brand.color_mid_gray)

        c1 = tcl.rows[ri].cells[1]
        _set_bg(c1, bg)
        _set_col_w(c1, 7760)
        _set_padding(c1, top=60, bottom=60, left=120)
        _spacing(c1.paragraphs[0], before=0, after=0)
        _run(c1.paragraphs[0], val, brand, bold=True, size=9.5,
             color=brand.color_accent_1)

    _divider(doc, _rgb_hex(brand.color_accent_2), sz=6, before=60, after=60)

    # ── 3. FECHAS DEL SERVICIO ────────────────────────────────────────────────
    if quote.fechas:
        _section_label(doc, brand, "Fechas del servicio")
        _build_fechas_table(doc, brand, quote.fechas)

    _divider(doc, _rgb_hex(brand.color_accent_2), sz=6, before=70, after=60)

    # ── 4. TABLA DE EQUIPOS ───────────────────────────────────────────────────
    _section_label(doc, brand, "Equipos  —  precio por fecha")
    _build_items_table(doc, brand, quote.items, subtotal_fecha, n_fechas, total_general)

    _divider(doc, _rgb_hex(brand.color_accent_2), sz=6, before=70, after=60)

    # ── 5. BLOQUE INVERSIÓN ───────────────────────────────────────────────────
    _section_label(doc, brand, "Inversión")

    WI0, WI1 = 3800, 5560
    tinv = doc.add_table(rows=1, cols=2)
    tinv.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(tinv, FULL)
    _no_borders(tinv)

    ci = tinv.rows[0].cells[0]
    _set_bg(ci, brand.color_primary_dark)
    _set_col_w(ci, WI0)
    _set_padding(ci, top=120, bottom=120, left=120, right=120)
    ci.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = ci.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=0)
    _run(p, "TOTAL GENERAL\n", brand, bold=True, size=8, color=brand.color_mid_gray)
    _run(p, _cop(total_general), brand, bold=True, size=18, color=brand.color_accent_1)

    cd = tinv.rows[0].cells[1]
    _set_bg(cd, brand.color_light_bg)
    _set_col_w(cd, WI1)
    _set_padding(cd, top=100, bottom=100, left=140, right=100)
    cd.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cd.paragraphs[0]
    _spacing(p, before=0, after=18)
    _run(p, "Forma de pago\n", brand, bold=True, size=8.5, color=brand.color_primary_dark)
    _run(p, quote.forma_pago, brand, size=8.5, color=brand.color_body_text)
    if quote.nota_pago:
        _run(p, f"\n{quote.nota_pago}", brand, italic=True, size=7.5,
             color=brand.color_mid_gray)

    incluye_lines = f"· {subtotal_fecha:,.0f}".replace(",", ".") + " por fecha\n"
    incluye_lines += f"· {n_fechas} fecha{'s' if n_fechas != 1 else ''} de evento\n"
    for item in quote.items:
        incluye_lines += f"· {item.descripcion}\n"
    if quote.incluye_extra:
        incluye_lines += quote.incluye_extra

    p2 = cd.add_paragraph()
    _spacing(p2, before=30, after=0)
    _run(p2, "Incluye\n", brand, bold=True, size=8.5, color=brand.color_primary_dark)
    _run(p2, incluye_lines.rstrip(), brand, size=8, color=brand.color_body_text)

    _divider(doc, _rgb_hex(brand.color_accent_2), sz=6, before=70, after=60)

    # ── 6. VIGENCIA + FIRMA ───────────────────────────────────────────────────
    WV = 4680
    tvf = doc.add_table(rows=1, cols=2)
    tvf.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(tvf, FULL)
    _no_borders(tvf)

    cv = tvf.rows[0].cells[0]
    _set_bg(cv, brand.color_light_bg)
    _set_col_w(cv, WV)
    _set_padding(cv)
    p = cv.paragraphs[0]
    _spacing(p, before=0, after=16)
    _run(p, "VIGENCIA DE LA COTIZACIÓN\n", brand, bold=True, size=8.5,
         color=brand.color_primary_dark)
    _run(p, f"Emisión:        {fecha_emision}\n", brand, size=8.5,
         color=brand.color_body_text)
    _run(p, f"Válida hasta:  {fecha_vigencia}", brand, size=8.5,
         color=brand.color_body_text)

    cf = tvf.rows[0].cells[1]
    _set_bg(cf, brand.color_white)
    _set_col_w(cf, WV)
    _set_padding(cf)
    p = cf.paragraphs[0]
    _spacing(p, before=0, after=16)
    _run(p, "PROPONENTE\n", brand, bold=True, size=8.5, color=brand.color_primary_dark)
    _run(p, f"{brand.proponent_name}\n", brand, bold=True, size=9,
         color=brand.color_body_text)
    _run(p, f"{brand.proponent_id_full}\n\n", brand, size=8.5,
         color=brand.color_body_text)
    _run(p, "___________________________\nFirma", brand, size=8,
         color=brand.color_mid_gray)

    # ── 7. PIE ────────────────────────────────────────────────────────────────
    _divider(doc, _rgb_hex(brand.color_accent_1), sz=6, before=60, after=20)
    footer = brand.proponent_full
    if brand.proponent_id_full:
        footer += f"  ·  {brand.proponent_id_full}"
    _para(doc, brand, footer, size=7.5, color=brand.color_mid_gray,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    # ── 8. BANNER ─────────────────────────────────────────────────────────────
    if brand.banner_path:
        pb = doc.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(pb, before=40, after=0)
        pb.add_run().add_picture(brand.banner_path, width=Cm(17))

    doc.save(output_path)
    print(f"Cotizacion generada: {output_path}")
    return output_path
