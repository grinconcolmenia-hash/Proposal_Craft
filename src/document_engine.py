"""
ProposalCraft — Motor Universal de Documentos
El agente (Claude) decide qué secciones necesita el documento y en qué orden.
Este motor renderiza lo que el agente especificó — sin estructura fija.

Uso:
    from src.brand_loader import load_brand
    from src.document_engine import (
        generate_document, DocumentSpec,
        HeaderSection, TextSection, TableSection,
        InversionSection, VigenciaFirmaSection,
        Column, Row,
    )

    brand = load_brand()
    spec  = DocumentSpec(
        output_filename = "Cotizacion_Cliente.docx",
        sections = [
            HeaderSection(
                cliente    = "UNEQUAL",
                nit_cliente= "NIT 901170356-3",
                subtitulo  = "Copa Mundial de Futbol - Grupo K",
            ),
            TableSection(
                titulo  = "Fechas del servicio",
                columns = [
                    Column("#",       0.10, "center"),
                    Column("Partido", 0.55, "left"),
                    Column("Fecha / Hora", 0.35, "center"),
                ],
                rows = [
                    Row(["Partido 1", "Uzbekistan vs Colombia", "17 jun - 9:00 p.m."]),
                    Row(["Partido 2", "Colombia vs RD Congo",   "23 jun - 9:00 p.m."], "alt"),
                    Row(["Partido 3", "Colombia vs Portugal",   "27 jun - 6:30 p.m."]),
                ],
            ),
            TableSection(
                titulo  = "Equipos  -  precio por fecha",
                columns = [
                    Column("N°",          0.07, "center"),
                    Column("Descripcion", 0.73, "left"),
                    Column("Valor",       0.20, "right"),
                ],
                rows = [
                    Row(["01", "Pantalla 3x5",                   "$ 3.750.000"]),
                    Row(["02", "4 QSC relevos",                  "$   800.000"], "alt"),
                    Row(["03", "2 Electro Voice PA",             "$   400.000"]),
                    Row(["04", "2 Microfonos",                   "$   200.000"], "alt"),
                    Row(["05", "4 Cabezas beam 230 por DMX",     "$   800.000"]),
                    Row(["06", "20 Par leds perimetral por DMX", "$   600.000"], "alt"),
                    Row(["",   "SUBTOTAL POR FECHA",             "$ 6.550.000"], "subtotal"),
                    Row(["",   "TOTAL 3 FECHAS",                 "$ 19.650.000"], "total"),
                ],
            ),
            InversionSection(
                valor_total = "$ 19.650.000",
                forma_pago  = "Pago a 30 dias",
                nota        = "Precio aplica por cada fecha de evento.",
                incluye     = (
                    "- Pantalla 3x5\\n"
                    "- 4 QSC relevos\\n"
                    "- 2 Electro Voice PA\\n"
                    "- 2 Microfonos\\n"
                    "- 4 Cabezas beam 230 DMX\\n"
                    "- 20 Par leds perimetral DMX"
                ),
            ),
            VigenciaFirmaSection(),
        ],
    )
    path = generate_document(brand, spec)
"""

import os
from dataclasses import dataclass, field
from datetime import date, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.brand_loader import BrandConfig


# ══════════════════════════════════════════════════════════════════════════════
# ESTRUCTURAS DE DATOS — el agente construye estos objetos
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class HeaderSection:
    """Bloque de encabezado: logo + datos del cliente."""
    cliente: str
    subtitulo: str
    nit_cliente: str = ""   # opcional


@dataclass
class TextSection:
    """Párrafo de texto con título de sección."""
    titulo: str             # ej: "RESUMEN EJECUTIVO"
    cuerpo: str             # texto libre, soporta \n para saltos de línea


@dataclass
class Column:
    """Define una columna de tabla."""
    label: str              # texto del encabezado de columna
    width_pct: float        # proporción del ancho total (todas deben sumar 1.0)
    align: str = "left"     # "left" | "center" | "right"


@dataclass
class Row:
    """
    Una fila de tabla. Las celdas son strings.

    Estilos disponibles:
      normal   → fondo light_bg
      alt      → fondo white (alternancia con normal)
      dark     → fondo primary_dark, texto accent_1  (header de grupo)
      subtotal → fondo accent_2, texto accent_1, bold
      total    → fondo primary_dark, texto accent_1, bold grande
    """
    cells: list
    style: str = "normal"


@dataclass
class TableSection:
    """Tabla completamente libre: el agente define columnas y filas."""
    titulo: str             # label de sección encima de la tabla
    columns: list           # list[Column]
    rows: list              # list[Row]


@dataclass
class InversionSection:
    """Bloque visual de inversión / total + condiciones de pago."""
    valor_total: str        # ej: "$ 19.650.000"
    forma_pago: str
    nota: str = ""          # texto pequeño bajo forma de pago (italic)
    incluye: str = ""       # texto bajo "Incluye" (usa \n para ítems)


@dataclass
class VigenciaFirmaSection:
    """Bloque de vigencia de la propuesta y firma del proponente."""
    fecha_emision: str = ""     # auto-calculado si vacío
    fecha_vigencia: str = ""    # auto-calculado si vacío


@dataclass
class DocumentSpec:
    """Especificación completa de un documento. El agente la construye."""
    sections: list              # list[Section] — cualquier combinación y orden
    output_filename: str = ""   # si vacío → "Documento.docx"


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE BAJO NIVEL (sin hardcoding de marca)
# ══════════════════════════════════════════════════════════════════════════════

FULL_TWIPS = 9360   # ancho útil en twips (página carta, márgenes 2.54 cm)

_ALIGN = {
    "left":   WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right":  WD_ALIGN_PARAGRAPH.RIGHT,
}


def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    tcPr.append(shd)


def _set_padding(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_col_w(cell, twips: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _tbl_width(table, twips: int):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    bdr = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        bdr.append(b)
    tblPr.append(bdr)


def _spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    if line:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _run(para, text: str, brand: BrandConfig,
         bold=False, italic=False, size=9.5, color: RGBColor = None):
    if color is None:
        color = brand.color_body_text
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = brand.font
    r.font.size = Pt(size)
    r.font.color.rgb = color


def _para(doc, brand: BrandConfig, text="", bold=False, italic=False,
          size=9.5, color: RGBColor = None,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          before=0, after=60, line=None):
    if color is None:
        color = brand.color_body_text
    p = doc.add_paragraph()
    p.alignment = align
    _spacing(p, before=before, after=after, line=line)
    if text:
        _run(p, text, brand, bold=bold, italic=italic, size=size, color=color)
    return p


def _divider(doc, color_hex: str, sz=6, before=70, after=60):
    p = doc.add_paragraph()
    _spacing(p, before=before, after=after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _section_label(doc, brand: BrandConfig, text: str):
    """Etiqueta de sección en accent_1, sobre la tabla o bloque."""
    _para(doc, brand, text.upper(), bold=True, size=8,
          color=brand.color_accent_1, before=0, after=24)


# ══════════════════════════════════════════════════════════════════════════════
# RENDERERS POR TIPO DE SECCIÓN
# ══════════════════════════════════════════════════════════════════════════════

def _render_header(doc, brand: BrandConfig, s: HeaderSection):
    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_logo, before=0, after=40)
    logo_path = brand.logo_dark_path or brand.logo_light_path
    if logo_path:
        p_logo.add_run().add_picture(logo_path, width=Cm(4))
    else:
        _run(p_logo, brand.company_name, brand,
             bold=True, size=16, color=brand.color_accent_1)

    _divider(doc, _rgb_hex(brand.color_accent_1), sz=12, before=20, after=50)

    # Etiqueta COTIZACIÓN / PROPUESTA COMERCIAL
    _para(doc, brand, "PROPUESTA COMERCIAL", bold=True, size=7.5,
          color=brand.color_mid_gray, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=0, after=4)

    if s.subtitulo:
        _para(doc, brand, s.subtitulo, italic=True, size=9,
              color=brand.color_mid_gray, align=WD_ALIGN_PARAGRAPH.CENTER,
              before=0, after=14)

    # Bloque cliente
    rows_cliente = [("Cliente", s.cliente)]
    if s.nit_cliente:
        rows_cliente.append(("NIT / ID", s.nit_cliente))

    tcl = doc.add_table(rows=len(rows_cliente), cols=2)
    tcl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_width(tcl, FULL_TWIPS)
    _no_borders(tcl)

    for ri, (lbl, val) in enumerate(rows_cliente):
        c0 = tcl.rows[ri].cells[0]
        _set_bg(c0, brand.color_primary_dark)
        _set_col_w(c0, 1600)
        _set_padding(c0, top=60, bottom=60, left=100, right=80)
        _spacing(c0.paragraphs[0], before=0, after=0)
        _run(c0.paragraphs[0], lbl, brand, bold=True, size=8,
             color=brand.color_mid_gray)

        c1 = tcl.rows[ri].cells[1]
        _set_bg(c1, brand.color_primary_dark)
        _set_col_w(c1, 7760)
        _set_padding(c1, top=60, bottom=60, left=120, right=80)
        _spacing(c1.paragraphs[0], before=0, after=0)
        _run(c1.paragraphs[0], val, brand, bold=True, size=9.5,
             color=brand.color_accent_1)


def _render_text(doc, brand: BrandConfig, s: TextSection):
    _section_label(doc, brand, s.titulo)
    _para(doc, brand, s.cuerpo, size=9.5, color=brand.color_body_text,
          before=0, after=0, line=264)


def _render_table(doc, brand: BrandConfig, s: TableSection):
    _section_label(doc, brand, s.titulo)

    n_cols = len(s.columns)
    n_rows = 1 + len(s.rows)   # header + data

    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(t, FULL_TWIPS)
    _no_borders(t)

    # Anchos en twips
    widths = [int(FULL_TWIPS * col.width_pct) for col in s.columns]

    # ── Header de tabla ───────────────────────────────────────────────────────
    for ci, (col, w) in enumerate(zip(s.columns, widths)):
        c = t.rows[0].cells[ci]
        _set_bg(c, brand.color_primary_dark)
        _set_col_w(c, w)
        _set_padding(c, top=70, bottom=70)
        p = c.paragraphs[0]
        p.alignment = _ALIGN.get(col.align, WD_ALIGN_PARAGRAPH.LEFT)
        _spacing(p, before=0, after=0)
        if col.label:
            _run(p, col.label, brand, bold=True, size=7.5,
                 color=brand.color_accent_1)

    # ── Filas de datos ────────────────────────────────────────────────────────
    for ri, row in enumerate(s.rows):
        trow = t.rows[ri + 1]

        # Paleta según estilo de fila
        style = (row.style or "normal").lower()
        if style == "total":
            bg        = brand.color_primary_dark
            fg        = brand.color_accent_1
            txt_bold  = True
            txt_size  = 11
            pad_v     = 110
        elif style == "subtotal":
            bg        = brand.color_accent_2
            fg        = brand.color_accent_1
            txt_bold  = True
            txt_size  = 9
            pad_v     = 90
        elif style == "dark":
            bg        = brand.color_primary_dark
            fg        = brand.color_accent_1
            txt_bold  = True
            txt_size  = 8.5
            pad_v     = 80
        elif style == "alt":
            bg        = brand.color_white
            fg        = brand.color_body_text
            txt_bold  = False
            txt_size  = 9.5
            pad_v     = 80
        else:   # normal
            bg        = brand.color_light_bg
            fg        = brand.color_body_text
            txt_bold  = False
            txt_size  = 9.5
            pad_v     = 80

        for ci, (col, w) in enumerate(zip(s.columns, widths)):
            cell = trow.cells[ci]
            _set_bg(cell, bg)
            _set_col_w(cell, w)
            _set_padding(cell, top=pad_v, bottom=pad_v,
                         left=140 if col.align == "left" else 80,
                         right=140 if col.align == "right" else 80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = _ALIGN.get(col.align, WD_ALIGN_PARAGRAPH.LEFT)
            _spacing(p, before=0, after=0)

            text = row.cells[ci] if ci < len(row.cells) else ""
            if text:
                _run(p, str(text), brand, bold=txt_bold,
                     size=txt_size, color=fg)


def _render_inversion(doc, brand: BrandConfig, s: InversionSection):
    _section_label(doc, brand, "Inversion")

    WI0, WI1 = 3800, 5560
    tinv = doc.add_table(rows=1, cols=2)
    tinv.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(tinv, FULL_TWIPS)
    _no_borders(tinv)

    # Celda izquierda — total
    ci = tinv.rows[0].cells[0]
    _set_bg(ci, brand.color_primary_dark)
    _set_col_w(ci, WI0)
    _set_padding(ci, top=130, bottom=130, left=120, right=120)
    ci.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = ci.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=0)
    _run(p, "TOTAL\n", brand, bold=True, size=8, color=brand.color_mid_gray)
    _run(p, s.valor_total, brand, bold=True, size=18, color=brand.color_accent_1)

    # Celda derecha — condiciones
    cd = tinv.rows[0].cells[1]
    _set_bg(cd, brand.color_light_bg)
    _set_col_w(cd, WI1)
    _set_padding(cd, top=110, bottom=110, left=140, right=100)
    cd.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = cd.paragraphs[0]
    _spacing(p, before=0, after=16)
    _run(p, "Forma de pago\n", brand, bold=True, size=8.5,
         color=brand.color_primary_dark)
    _run(p, s.forma_pago, brand, size=8.5, color=brand.color_body_text)
    if s.nota:
        _run(p, f"\n{s.nota}", brand, italic=True, size=7.5,
             color=brand.color_mid_gray)

    if s.incluye:
        p2 = cd.add_paragraph()
        _spacing(p2, before=30, after=0)
        _run(p2, "Incluye\n", brand, bold=True, size=8.5,
             color=brand.color_primary_dark)
        _run(p2, s.incluye, brand, size=8, color=brand.color_body_text)


def _render_vigencia_firma(doc, brand: BrandConfig, s: VigenciaFirmaSection):
    MESES = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]

    def fmt(d: date) -> str:
        return f"{d.day} de {MESES[d.month - 1]} de {d.year}"

    today = date.today()
    emision  = s.fecha_emision  or fmt(today)
    vigencia = s.fecha_vigencia or fmt(today + timedelta(days=brand.validity_days))

    WV = 4680
    tvf = doc.add_table(rows=1, cols=2)
    tvf.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(tvf, FULL_TWIPS)
    _no_borders(tvf)

    cv = tvf.rows[0].cells[0]
    _set_bg(cv, brand.color_light_bg)
    _set_col_w(cv, WV)
    _set_padding(cv)
    p = cv.paragraphs[0]
    _spacing(p, before=0, after=16)
    _run(p, "VIGENCIA\n", brand, bold=True, size=8.5,
         color=brand.color_primary_dark)
    _run(p, f"Emision:       {emision}\n", brand, size=8.5,
         color=brand.color_body_text)
    _run(p, f"Valida hasta:  {vigencia}", brand, size=8.5,
         color=brand.color_body_text)

    cf = tvf.rows[0].cells[1]
    _set_bg(cf, brand.color_white)
    _set_col_w(cf, WV)
    _set_padding(cf)
    p = cf.paragraphs[0]
    _spacing(p, before=0, after=16)
    _run(p, "PROPONENTE\n", brand, bold=True, size=8.5,
         color=brand.color_primary_dark)
    _run(p, f"{brand.proponent_name}\n", brand, bold=True, size=9,
         color=brand.color_body_text)
    _run(p, f"{brand.proponent_id_full}\n\n", brand, size=8.5,
         color=brand.color_body_text)
    _run(p, "___________________________\nFirma", brand, size=8,
         color=brand.color_mid_gray)


def _render_footer(doc, brand: BrandConfig):
    _divider(doc, _rgb_hex(brand.color_accent_1), sz=6, before=60, after=20)
    footer = brand.proponent_full
    if brand.proponent_id_full:
        footer += f"  -  {brand.proponent_id_full}"
    _para(doc, brand, footer, size=7.5, color=brand.color_mid_gray,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    if brand.banner_path:
        pb = doc.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(pb, before=40, after=0)
        pb.add_run().add_picture(brand.banner_path, width=Cm(17))


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

# Mapa de tipo de sección → renderer
_RENDERERS = {
    HeaderSection:       _render_header,
    TextSection:         _render_text,
    TableSection:        _render_table,
    InversionSection:    _render_inversion,
    VigenciaFirmaSection:_render_vigencia_firma,
}


def generate_document(brand: BrandConfig, spec: DocumentSpec,
                      output_path: str | None = None) -> str:
    """
    Genera un .docx a partir de un DocumentSpec.

    Args:
        brand:       Configuración de marca (cargada con load_brand()).
        spec:        Especificación del documento construida por el agente.
        output_path: Ruta de salida opcional. Si None usa brand.output_dir.

    Returns:
        Ruta absoluta del archivo generado.
    """
    if output_path is None:
        fname = spec.output_filename or "Documento.docx"
        output_path = os.path.join(brand.output_dir, fname)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    for i, section in enumerate(spec.sections):
        # Divisor entre secciones (salvo antes del header)
        if i > 0 and not isinstance(section, VigenciaFirmaSection):
            _divider(doc, _rgb_hex(brand.color_accent_2), sz=6,
                     before=70, after=60)

        renderer = _RENDERERS.get(type(section))
        if renderer:
            renderer(doc, brand, section)
        else:
            raise TypeError(f"Tipo de sección no reconocido: {type(section)}")

    # Footer siempre al final
    _render_footer(doc, brand)

    doc.save(output_path)
    print(f"Documento generado: {output_path}")
    return output_path
